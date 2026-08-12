from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

out_dir = Path("samples")
out_dir.mkdir(exist_ok=True)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title_run = title.add_run("Catatan Psikologi Penerbang - Format Naratif Tidak Seragam")
title_run.font.name = "Arial"
title_run.font.size = Pt(18)
title_run.bold = True

doc.add_paragraph(
    "Dokumen ini sengaja berisi paragraf naratif, label campur, singkatan, dan tabel kecil agar fitur Import Cerdas CSAKT dapat menguji entity extraction."
)
doc.add_paragraph(
    "NRP: 522884, nama Letkol Pnb Chandra H.; pemeriksaan 21 April 2024 menunjukkan stress index 54, atensi 71, stabilitas emosi 67. Rekomendasi psikolog: pendampingan psikologi operasional."
)
doc.add_paragraph(
    "Kapten Pnb Bagas R / NRP 531115 menjalani evaluasi tanggal 03/06/2025. Skor kognitif baik, stres 46, beban kognitif 48, rekomendasi: monitoring workload."
)
doc.add_paragraph(
    "Catatan bebas: Mayor Galih N. punya stress tinggi, NRP tidak ditulis pada paragraf ini sehingga harus ditandai sebagai error atau perlu konfirmasi."
)

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
headers = ["Personel", "tgl psikotes", "stres", "atensi", "saran psikolog"]
for idx, header in enumerate(headers):
    table.rows[0].cells[idx].text = header

rows = [
    ["529701 - Aditya W.", "12 Maret 2016", "22", "89", "Sesuai profil tugas"],
    ["527334 - Galih N.", "not a date", "68", "62", "Pendampingan dan review lanjutan"],
]
for row in rows:
    cells = table.add_row().cells
    for idx, value in enumerate(row):
        cells[idx].text = value

doc.add_paragraph("Footer manual: subtotal, tanda tangan, dan catatan kaki tidak perlu diimpor.")
doc.save(out_dir / "catatan_psikologi_naratif.docx")
print("sample docx created")
